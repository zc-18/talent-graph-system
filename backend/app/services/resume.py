"""简历解析（PDF / Word / 纯文本）+ 技能要素抽取。

赛题指标：简历技能提取准确率 ≥ 90%。
策略：大模型结构化抽取为主 + 词典规则兜底，技能统一归一化。
"""
from __future__ import annotations
import io
import re
from dataclasses import dataclass
from .. import clients
from .taxonomy import normalize_skill, skill_category, skill_type, SYNONYMS


@dataclass
class ResumeFileError(Exception):
    code: str
    message: str
    status_code: int = 422

    def __str__(self) -> str:
        return self.message


def extract_text(filename: str, content: bytes) -> str:
    """Extract PDF/DOCX/TXT with stable, machine-readable failure codes."""
    name = (filename or "").lower()
    if not content:
        raise ResumeFileError("EMPTY_FILE", "文件为空")
    if name.endswith(".pdf"):
        return _from_pdf(content)
    if name.endswith(".docx"):
        return _from_docx(content)
    if name.endswith(".doc"):
        raise ResumeFileError("UNSUPPORTED_DOC", "旧版 .doc 不受支持，请转换为 DOCX", 415)
    if not name.endswith(".txt"):
        raise ResumeFileError("UNSUPPORTED_FILE_TYPE", "仅支持 PDF、DOCX、TXT", 415)
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ResumeFileError("CORRUPT_FILE", "TXT 不是有效的 UTF-8 文本") from exc
    if not text.strip():
        raise ResumeFileError("EMPTY_DOCUMENT", "文档没有可解析文本")
    return text


def _from_pdf(content: bytes) -> str:
    # Standard-encrypted PDFs carry an /Encrypt entry. Some pdfminer versions
    # fail earlier with a generic parse error, so detect this stable marker first.
    if b"/Encrypt" in content:
        raise ResumeFileError("ENCRYPTED_FILE", "PDF 已加密，无法解析")
    try:
        import pdfplumber
        text = []
        has_images = False
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
                has_images = has_images or bool(page.images)
        result = "\n".join(text)
        if not result.strip():
            code = "SCANNED_PDF" if has_images else "EMPTY_DOCUMENT"
            message = "扫描版 PDF 需要 OCR" if has_images else "PDF 没有可解析文本"
            raise ResumeFileError(code, message)
        return result
    except ResumeFileError:
        raise
    except Exception as exc:  # noqa: BLE001
        marker = f"{type(exc).__name__} {exc}".lower()
        if "password" in marker or "encrypt" in marker:
            raise ResumeFileError("ENCRYPTED_FILE", "PDF 已加密，无法解析") from exc
        raise ResumeFileError("CORRUPT_FILE", "PDF 文件损坏或格式无效") from exc


def _from_docx(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(c.text for c in row.cells))
        result = "\n".join(parts)
        if not result.strip():
            raise ResumeFileError("EMPTY_DOCUMENT", "DOCX 没有可解析文本")
        return result
    except ResumeFileError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ResumeFileError("CORRUPT_FILE", "DOCX 文件损坏或格式无效") from exc


_RESUME_SYS = """你是简历解析专家。从简历文本中精确抽取候选人信息，只抽取简历中真实出现的内容，不臆造。
技能要细粒度到具体技术点（如 PyTorch、Spark、向量数据库、模型微调）。只输出JSON。"""

_RESUME_TPL = """解析以下简历，输出JSON：
{{
  "candidate_name": "姓名(无则空)",
  "years_experience": 工作年限数字,
  "education": "最高学历",
  "skills": ["技能点1","技能点2"],
  "skill_levels": {{"技能点":"familiar/proficient/expert"}},
  "projects": ["项目简述"],
  "titles": ["曾任岗位"]
}}

简历文本：
---
{resume}
---"""


def parse_resume(text: str) -> dict:
    """结构化解析简历。"""
    if not text or len(text.strip()) < 10:
        return {"candidate_name": "", "years_experience": 0, "skills": [], "skill_levels": {},
                "education": "", "projects": [], "titles": [], "raw_skill_count": 0,
                "raw_skill_terms": []}
    messages = [
        {"role": "system", "content": _RESUME_SYS},
        {"role": "user", "content": _RESUME_TPL.format(resume=text[:5000])},
    ]
    data = clients.chat_json(messages, temperature=0.1, max_tokens=1200)
    return _postprocess_resume(data, text)


def _postprocess_resume(data: dict, text: str) -> dict:
    raw_skills = data.get("skills", []) or []
    levels_in = data.get("skill_levels", {}) or {}
    norm_levels, norm_skills, seen = {}, [], set()
    raw_terms: list[str] = []                # 归一化**前**的原始表述（供别名学习，意见⑧）
    for s in raw_skills:
        term = s if isinstance(s, str) else s.get("name", "")
        if isinstance(term, str) and term.strip():
            raw_terms.append(term.strip())
        nm = normalize_skill(term)
        if not nm or nm in seen:
            continue
        seen.add(nm)
        norm_skills.append(nm)
        lvl = levels_in.get(s) if isinstance(s, str) else None
        norm_levels[nm] = lvl if lvl in ("familiar", "proficient", "expert") else "proficient"
    # 规则兜底：补充词典命中但模型漏抽的技能
    for kw, nm in SYNONYMS.items():
        if nm in seen:
            continue
        if re.search(re.escape(kw), text, re.IGNORECASE):
            seen.add(nm)
            norm_skills.append(nm)
            norm_levels[nm] = "familiar"
    try:
        yrs = float(data.get("years_experience", 0) or 0)
    except (ValueError, TypeError):
        yrs = 0.0
    return {
        "candidate_name": (data.get("candidate_name") or "").strip(),
        "years_experience": yrs,
        "education": data.get("education", ""),
        "skills": norm_skills,
        "skill_levels": norm_levels,
        "skill_categories": {s: skill_category(s) for s in norm_skills},
        "projects": data.get("projects", [])[:10],
        "titles": data.get("titles", [])[:5],
        "raw_skill_count": len(raw_skills),
        "raw_skill_terms": raw_terms,
    }


# 个人信息（PII）字段：依据数据合规与隐私最小化原则，不在服务端持久化
_PII_FIELDS = {"candidate_name", "projects", "titles"}


# ---------------- 简历正文脱敏（落盘/入库前的强制关口，意见⑧）----------------
# collect/base.py::mask_pii 只覆盖中国大陆手机号/邮箱/微信，简历语料还会出现
# 美式电话、QQ、身份证、个人主页链接，故在此扩一层。两者叠加使用。
_RE_PHONE_CN = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")
_RE_PHONE_INTL = re.compile(r"(?<![\d-])(?:\+?\d{1,2}[ .-]?)?\(?\d{3}\)?[ .-]\d{3}[ .-]\d{4}(?![\d-])")
_RE_EMAIL = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_RE_WECHAT = re.compile(r"(微信|weixin|wechat)\s*[:：]?\s*[A-Za-z0-9_-]{5,20}", re.I)
_RE_QQ = re.compile(r"(QQ|qq)\s*[:：]?\s*\d{5,12}")
_RE_IDCARD = re.compile(r"(?<!\d)\d{17}[\dXx](?!\d)")
_RE_URL = re.compile(r"https?://[^\s，,；;）)】\]]+")


def mask_contacts(text: str) -> str:
    """抹掉简历正文里的联系方式与个人链接。

    简历不同于 JD：JD 是企业主动公开的商业信息，正文原样留档；简历涉及真人，
    因此归档前先过这道关口。**注意口径**：这里去掉的是**联系方式**，不是"去标识化"——
    姓名/年龄/性别不在本函数处理范围内，归档 jsonl 仍属可识别个人数据，
    只有入库形态 `TalentProfile` 才是结构上不含身份的
    （见 采集合规说明「简历语料来源与合规」的分层口径说明）。
    """
    if not text:
        return text
    text = _RE_EMAIL.sub("[邮箱已脱敏]", text)
    text = _RE_PHONE_CN.sub("[电话已脱敏]", text)
    text = _RE_PHONE_INTL.sub("[电话已脱敏]", text)
    text = _RE_WECHAT.sub(r"\1:[已脱敏]", text)
    text = _RE_QQ.sub(r"\1:[已脱敏]", text)
    text = _RE_IDCARD.sub("[证件号已脱敏]", text)
    text = _RE_URL.sub("[链接已脱敏]", text)
    return text


def contains_contacts(text: str) -> bool:
    """脱敏结果自检：还能匹配到联系方式就说明漏了（供落盘前断言用）。"""
    return any(p.search(text or "") for p in
               (_RE_EMAIL, _RE_PHONE_CN, _RE_PHONE_INTL, _RE_QQ, _RE_IDCARD, _RE_URL))


def redact_for_storage(parsed: dict) -> dict:
    """数据最小化：剔除姓名/项目经历/任职单位等个人身份信息(PII)，
    仅保留用于岗位匹配分析的非身份技能要素，供（可选的）服务端留存。

    原始简历全文与姓名仅在内存中用于本次解析、即时返回给本人，绝不落库。
    """
    return {k: v for k, v in (parsed or {}).items() if k not in _PII_FIELDS}
