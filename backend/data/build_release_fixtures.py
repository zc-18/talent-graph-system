"""Build deterministic, source-grounded engineering regression fixtures.

The source corpora are already archived in this repository. This script does not
download data or call a model. Labels come from source-provided ``skill_tags`` or
``Skill Details`` fields. They are independent from TalentGraph output, but are
not dual-human-annotated release truth. PDF/DOCX files are generated from source
text and must not be counted as original resume files. Contact details are masked.
"""
from __future__ import annotations

import hashlib
import json
import re
import sys
import textwrap
import unicodedata
from collections import defaultdict, deque
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw
from pypdf import PdfReader, PdfWriter

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.services.resume import mask_contacts
from app.services.job_resolution import resolve_job_query
from app.services.taxonomy import SKILL_CATEGORY, SOFT_SKILLS, SYNONYMS, normalize_skill


HERE = Path(__file__).resolve().parent
OUT = HERE / "eval_fixtures"
FILES = OUT / "files"
JD_SOURCE = HERE / "raw" / "2024hist-r1" / "dataset_aijob2024.jsonl"
RESUME_SOURCE = HERE / "resumes" / "_datasets" / "hf_brackozi_Resume.jsonl"
SOURCE_URL = "https://huggingface.co/datasets/brackozi/Resume"

_SKILL_LINE = re.compile(r"(?im)^\s*([^\n-]{1,80})-\s*Exprience\s*-")
_BARE_PROFILE = re.compile(r"\b(?:linkedin|github)\.com/\S+", re.I)
_CAMPUS_TERMS = ("应届", "校招", "校园招聘", "毕业生", "在校生", "实习生")
_JUNIOR_TERMS = ("初级", "助理", "专员")
_SENIOR_TERMS = ("高级", "资深", "专家", "负责人", "主任", "总监", "架构师")
_YEARS = re.compile(r"(?<!\d)(\d{1,2})\s*(?:年|年以上|年及以上)")
_NON_SKILL_TAGS = {
    "本科", "硕士", "不限", "其他", "计算机相关专业", "数学/统计相关专业",
    "团队管理经验", "优秀开源项目经历", "参加算法相关竞赛/获奖", "发表算法相关优秀论文",
    "算法研究", "算法开发", "算法设计", "算法基础", "研究导向",
}
_ALIASES: dict[str, set[str]] = defaultdict(set)
for _alias, _canonical in SYNONYMS.items():
    _ALIASES[_canonical].add(_alias)

ROLE_CONTRACTS = {
    # Frozen contracts were curated from role definitions before this holdout was sampled.
    "Data Science": ("数据科学工程师", ["Python", "机器学习", "TensorFlow", "SQL", "数据挖掘"]),
    "Python Developer": ("Python开发工程师", ["Python", "MySQL", "JavaScript", "Linux", "Git"]),
    "Java Developer": ("Java开发工程师", ["Java", "Spring", "MySQL", "JavaScript", "Git"]),
    "DevOps Engineer": ("DevOps工程师", ["Python", "SQL", "Linux", "Git", "Docker"]),
    "Hadoop": ("大数据开发工程师", ["Hadoop", "Hive", "Python", "SQL", "Kafka"]),
    "ETL Developer": ("ETL开发工程师", ["ETL", "SQL", "数据仓库", "Python", "Git"]),
    "Database": ("数据库工程师", ["SQL", "MySQL", "Linux", "Redis", "Git"]),
}


def _jsonl(path: Path):
    with path.open(encoding="utf-8") as handle:
        for row, line in enumerate(handle, 1):
            if line.strip():
                yield row, json.loads(line)


def _skill(value: str) -> str:
    value = re.sub(r"\s+", " ", (value or "").strip(" *:;,.\t"))
    if not value or value in _NON_SKILL_TAGS or len(value) > 40:
        return ""
    return normalize_skill(value)


def _dedup_skills(values) -> list[str]:
    out: list[str] = []
    for value in values:
        skill = _skill(value)
        if skill and skill not in out:
            out.append(skill)
    return out


def _round_robin(groups: dict[str, list[dict]], count: int) -> list[dict]:
    queues = {key: deque(value) for key, value in sorted(groups.items()) if value}
    selected: list[dict] = []
    while queues and len(selected) < count:
        for key in list(queues):
            if queues[key]:
                selected.append(queues[key].popleft())
                if len(selected) == count:
                    break
            if not queues[key]:
                del queues[key]
    return selected


def _recruitment_type(title: str, text: str) -> str:
    blob = f"{title}\n{text}"
    if any(term in blob for term in _CAMPUS_TERMS):
        return "campus"
    if _YEARS.search(blob) or "社招" in blob or "社会招聘" in blob:
        return "social"
    return "unspecified"


def _seniority(title: str, text: str, recruitment_type: str) -> str:
    blob = f"{title}\n{text}"
    if any(term in title for term in _SENIOR_TERMS):
        return "senior"
    if recruitment_type == "campus" or any(term in title for term in _JUNIOR_TERMS):
        return "junior"
    years = [int(value) for value in _YEARS.findall(blob)]
    if years and max(years) >= 5:
        return "senior"
    if years and max(years) >= 2:
        return "middle"
    return "unspecified"


def _job_dimensions(title: str) -> dict[str, str]:
    resolution = resolve_job_query(title)
    return {
        "job": resolution.canonical_title or title or "unknown",
        "track": resolution.track,
        "industry": resolution.industry,
    }


def build_jd_holdout(count: int = 100) -> list[dict]:
    groups: dict[str, list[dict]] = defaultdict(list)
    for source_row, item in _jsonl(JD_SOURCE):
        extra = item.get("extra") or {}
        raw_text = item.get("raw_text") or ""
        folded_text = raw_text.casefold()
        tags = [skill for skill in _dedup_skills(
            re.split(r"[,，、;；|]", extra.get("skill_tags") or ""))
                if skill in SKILL_CATEGORY and (
                    skill.casefold() in folded_text or any(
                        alias.casefold() in folded_text for alias in _ALIASES.get(skill, ())))]
        if not tags:
            continue
        query = extra.get("query") or "unknown"
        job_title = (item.get("job_title") or "unknown").strip()
        recruitment_type = _recruitment_type(job_title, raw_text)
        seniority = _seniority(job_title, raw_text, recruitment_type)
        dimensions = _job_dimensions(job_title)
        groups["|".join((query, dimensions["job"], dimensions["track"],
                         dimensions["industry"], seniority, recruitment_type))].append({
            "id": f"real-jd-{source_row:04d}",
            "source_row": source_row,
            "source_url": item.get("url"),
            "domain": query,
            **dimensions,
            "job_title": job_title,
            "seniority": seniority,
            "recruitment_type": recruitment_type,
            "stratification_source": "source_title_and_explicit_text_markers",
            "raw_text": raw_text,
            "ground_truth_skills": tags,
            "annotation_source": "source_dataset_skill_tags",
            "annotation_complete": False,
            "truth_independent": False,
            "human_annotation_status": "not_performed",
        })
    selected = _round_robin(groups, count)
    if len(selected) < count:
        raise RuntimeError(f"only {len(selected)} JD fixtures available")
    return selected


def _resume_skills(text: str) -> list[str]:
    skills = [skill for skill in _dedup_skills(
        match.group(1) for match in _SKILL_LINE.finditer(text or ""))
              if skill in SKILL_CATEGORY or skill in SOFT_SKILLS]
    return skills[:30]


def _masked_resume(text: str) -> str:
    return _BARE_PROFILE.sub("[链接已脱敏]", mask_contacts(text or ""))


def _pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _ascii_lines(text: str, width: int = 94) -> list[str]:
    ascii_text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    lines: list[str] = []
    for paragraph in ascii_text.splitlines():
        lines.extend(textwrap.wrap(paragraph, width=width, replace_whitespace=False) or [""])
    return lines or [""]


def _pdf_bytes(text: str) -> bytes:
    pages = [list(chunk) for chunk in _chunks(_ascii_lines(text), 68)]
    page_ids = [3 + index * 2 for index in range(len(pages))]
    content_ids = [value + 1 for value in page_ids]
    font_id = 3 + len(pages) * 2
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{' '.join(f'{value} 0 R' for value in page_ids)}] /Count {len(pages)} >>".encode(),
    ]
    for page_id, content_id, lines in zip(page_ids, content_ids, pages):
        objects.append((f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                        f"/Resources << /Font << /F1 {font_id} 0 R >> >> "
                        f"/Contents {content_id} 0 R >>").encode())
        commands = ["BT", "/F1 8 Tf", "10 TL", "36 760 Td"]
        for line in lines:
            commands.extend([f"({_pdf_escape(line)}) Tj", "T*"])
        commands.append("ET")
        stream = "\n".join(commands).encode("latin-1")
        objects.append(f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    out = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_id, body in enumerate(objects, 1):
        offsets.append(len(out))
        out.extend(f"{object_id} 0 obj\n".encode() + body + b"\nendobj\n")
    xref = len(out)
    out.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    out.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        out.extend(f"{offset:010d} 00000 n \n".encode())
    out.extend((f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref}\n%%EOF\n").encode())
    return bytes(out)


def _chunks(values: list[str], size: int):
    for index in range(0, len(values), size):
        yield values[index:index + size]


def _write_docx(path: Path, text: str, two_column: bool) -> None:
    document = Document()
    document.core_properties.title = "De-identified resume evaluation fixture"
    if two_column:
        table = document.add_table(rows=1, cols=2)
        lines = text.splitlines()
        midpoint = max(1, len(lines) // 2)
        table.cell(0, 0).text = "\n".join(lines[:midpoint])
        table.cell(0, 1).text = "\n".join(lines[midpoint:])
    else:
        for paragraph in text.splitlines():
            document.add_paragraph(paragraph)
    document.save(path)


def build_resume_holdout(count: int = 120) -> list[dict]:
    groups: dict[str, list[dict]] = defaultdict(list)
    for source_row, item in _jsonl(RESUME_SOURCE):
        category = item.get("Category") or ""
        if category not in ROLE_CONTRACTS:
            continue
        raw_text = item.get("Resume") or ""
        skills = _resume_skills(raw_text)
        if len(skills) < 3:
            continue
        groups[category].append({"source_row": source_row, "category": category,
                                 "text": _masked_resume(raw_text), "skills": skills})
    selected = _round_robin(groups, count)
    if len(selected) < count:
        raise RuntimeError(f"only {len(selected)} resume fixtures available")

    formats = ["pdf", "docx", "txt"]
    rows: list[dict] = []
    for index, item in enumerate(selected, 1):
        file_format = formats[(index - 1) % len(formats)]
        filename = f"resume-{index:03d}.{file_format}"
        path = FILES / filename
        text = item.pop("text")
        if file_format == "pdf":
            path.write_bytes(_pdf_bytes(text))
        elif file_format == "docx":
            _write_docx(path, text, two_column=index % 2 == 0)
        else:
            path.write_text(text, encoding="utf-8")
        role, _ = ROLE_CONTRACTS[item["category"]]
        rows.append({
            "id": f"real-resume-{index:03d}",
            "source_row": item["source_row"],
            "source_url": SOURCE_URL,
            "license": "MIT",
            "category": item["category"],
            "target_job": role,
            "file": f"files/{filename}",
            "format": file_format,
            "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
            "ground_truth_skills": item["skills"],
            "annotation_source": "source_resume_skill_details",
            "annotation_complete": False,
            "truth_independent": False,
            "human_annotation_status": "not_performed",
            "file_origin": "generated_from_source_text",
            "is_original_source_file": False,
        })
    return rows


def build_match_pairs(resumes: list[dict], count: int = 120) -> list[dict]:
    pools: dict[str, list[tuple[dict, str]]] = {"high": [], "medium": [], "low": []}
    for resume in resumes:
        resume_skills = set(resume["ground_truth_skills"])
        for target_category, (_, capabilities) in ROLE_CONTRACTS.items():
            overlap = len(resume_skills & set(capabilities))
            if target_category == resume["category"] and overlap >= 2:
                pools["high"].append((resume, target_category))
            elif target_category != resume["category"] and overlap == 1:
                pools["medium"].append((resume, target_category))
            elif target_category != resume["category"] and overlap == 0:
                pools["low"].append((resume, target_category))
    per_label = count // 3
    selected = [(label, pair) for label in ("high", "medium", "low")
                for pair in pools[label][:per_label]]
    if any(len(pools[label]) < per_label for label in pools):
        raise RuntimeError("not enough balanced independent match pairs")
    pairs = []
    for index, (label, (resume, target_category)) in enumerate(selected, 1):
        role, capabilities = ROLE_CONTRACTS[target_category]
        dimensions = _job_dimensions(role)
        pairs.append({
            "id": f"match-{index:03d}", "resume_id": resume["id"],
            "target_job": role, "domain": target_category, **dimensions,
            "seniority": "unspecified", "recruitment_type": "unspecified",
            "target_category": target_category, "contract_capabilities": capabilities,
            "ground_truth_label": label,
            "annotation_source": "deterministic_category_and_skill_overlap_rules",
            "truth_independent": False,
            "human_annotation_status": "not_performed",
        })
    return pairs


def build_hr_batches(resumes: list[dict]) -> list[dict]:
    batches = []
    categories = list(ROLE_CONTRACTS)[:5]
    for batch_index, category in enumerate(categories, 1):
        role, capabilities = ROLE_CONTRACTS[category]
        dimensions = _job_dimensions(role)
        cap_set = set(capabilities)
        def overlap(row):
            return len(set(row["ground_truth_skills"]) & cap_set)
        relevant = sorted((row for row in resumes if row["category"] == category),
                          key=lambda row: (-overlap(row), row["id"]))[:8]
        distractors = sorted((row for row in resumes if row["category"] != category),
                             key=lambda row: (overlap(row), row["id"]))[:16]
        candidates = []
        for rank, resume in enumerate(relevant + distractors, 1):
            relevance = 2 if resume["category"] == category else (
                1 if set(resume["ground_truth_skills"]) & set(capabilities) else 0)
            candidates.append({"resume_id": resume["id"], "relevance": relevance,
                               "annotation_source": "deterministic_category_and_skill_overlap_rules"})
        batches.append({"id": f"hr-batch-{batch_index}", "target_job": role,
                        "domain": category, **dimensions,
                        "seniority": "unspecified", "recruitment_type": "unspecified",
                        "contract_capabilities": capabilities, "candidates": candidates,
                        "truth_independent": False,
                        "annotation_source": "deterministic_category_and_skill_overlap_rules",
                        "human_annotation_status": "not_performed"})
    return batches


def build_diagnostics() -> dict[str, str]:
    blank = FILES / "diagnostic-empty.pdf"
    blank.write_bytes(_pdf_bytes(""))
    scan = FILES / "diagnostic-scanned.pdf"
    image = Image.new("RGB", (1200, 1600), "white")
    ImageDraw.Draw(image).text((80, 100), "Scanned resume: Python, SQL, Docker", fill="black")
    image.save(scan, "PDF", resolution=150)
    legacy = FILES / "diagnostic-legacy.doc"
    legacy.write_bytes(b"Legacy Word binary placeholder")
    corrupt = FILES / "diagnostic-corrupt.pdf"
    corrupt.write_bytes(b"%PDF-1.4\ncorrupt")
    encrypted = FILES / "diagnostic-encrypted.pdf"
    reader = PdfReader(blank)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("talentgraph-test")
    with encrypted.open("wb") as handle:
        writer.write(handle)
    empty = FILES / "diagnostic-empty.txt"
    empty.write_bytes(b"")
    return {
        "scanned_pdf": f"files/{scan.name}", "legacy_doc": f"files/{legacy.name}",
        "encrypted": f"files/{encrypted.name}", "empty": f"files/{empty.name}",
        "corrupt": f"files/{corrupt.name}", "oversized": "generated-in-memory-by-evaluator",
    }


def _dump(name: str, value) -> None:
    (OUT / name).write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    FILES.mkdir(parents=True, exist_ok=True)
    jd = build_jd_holdout()
    resumes = build_resume_holdout()
    matches = build_match_pairs(resumes)
    hr_batches = build_hr_batches(resumes)
    _dump("jd_holdout.json", jd)
    _dump("resume_holdout.json", resumes)
    _dump("match_pairs.json", matches)
    _dump("hr_ranking_batches.json", hr_batches)
    _dump("diagnostics.json", build_diagnostics())
    _dump("fixture_manifest.json", {
        "jd_count": len(jd), "resume_count": len(resumes), "match_pair_count": len(matches),
        "hr_batch_count": len(hr_batches), "resume_source": SOURCE_URL, "resume_license": "MIT",
        "evidence_class": "engineering_regression_not_release_human_truth",
        "truth_policy": "source labels and deterministic rules; not dual-human-annotated",
        "jd_stratification_dimensions": [
            "domain", "job", "track", "industry", "seniority", "recruitment_type"],
        "matching_input_policy": "resume skills parsed from files; fixture truth skills are not injected",
        "jd_human_annotated_count": 0, "resume_original_file_count": 0,
        "resume_human_annotated_count": 0, "match_human_annotated_count": 0,
        "hr_human_labeled_candidate_count": 0,
    })
    print(f"release fixtures written to {OUT}")


if __name__ == "__main__":
    main()
