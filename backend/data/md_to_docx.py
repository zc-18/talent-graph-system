"""将提交用 Markdown 文档转换为 Word(.docx)。

支持：# / ## / ### / #### 标题、表格、- / · 列表、**加粗**、代码块、普通段落、引用。
"""
from __future__ import annotations
import os
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCS = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "docs"))
# Markdown 源文件统一放在 docs/_source/，交付目录 docs/ 下只保留生成的 .docx
SRC = os.path.join(DOCS, "_source")
FILES = ["作品设计与实现方案.md", "测试方案与报告.md", "部署说明.md", "演示视频脚本.md",
         "技术答辩文档.md", "改进说明_第二版.md", "改进说明_第三版.md"]

ACCENT = RGBColor(0x36, 0x52, 0xD9)
INK = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x55, 0x65, 0x73)


def _is_cjk(ch: str) -> bool:
    return bool(ch) and (
        "　" <= ch <= "〿"      # CJK 标点
        or "一" <= ch <= "鿿"   # 汉字
        or "＀" <= ch <= "￯"   # 全角
        or ch in "、。，；：？！（）《》「」【】—…·"
    )


def join_soft_wrap(parts):
    """把 Markdown 软换行的若干行拼回一个段落。

    必须先拼再做行内解析：源文件里有 `然后**直接对线上正在读的生产库\\n`talent_graph_v3`
    执行 apply**」` 这种**跨行加粗**，逐行解析时两个 `**` 各自落单、谁都配不上对，
    结果就是 Word 里原样打印出字面的星号（2026-08-30 在《改进说明》里实际踩到）。

    拼接规则：两侧都是中日韩字符/全角标点时不插空格（中文不靠空格断词），
    否则补一个空格，避免把英文单词粘成一个词。
    """
    out = ""
    for part in parts:
        part = part.rstrip()
        if not out:
            out = part
            continue
        if not part:
            continue
        if _is_cjk(out[-1]) and _is_cjk(part[0]):
            out += part
        else:
            out += " " + part
    return out


def _is_block_start(line: str) -> bool:
    """这一行是否开启了一个新块（因而不能被并进上一段）。"""
    s = line.strip()
    return (
        not s
        or s.startswith("#")
        or s.startswith(">")
        or s.startswith("|")
        or s.startswith("```")
        or s == "---"
        or bool(re.match(r"^[-*·]\s+", s))
        or bool(re.match(r"^\d+\.\s+", s))
    )


def add_runs(paragraph, text):
    """处理 **加粗** 与 `行内代码`，两者可互相嵌套。

    必须从左到右单遍扫描、谁先出现谁优先，不能先切一种再切另一种：
    - 先切加粗，`` `**已有**` `` 这种「用代码块展示字面星号」的写法会被拆坏；
    - 先切代码，``**…`parent_id`…**`` 这种加粗里套代码的长句会露出字面反引号。
    两个坑在《改进说明_第三版》和《作品设计与实现方案》里各中过一次。
    Markdown 规定行内代码优先级高于强调，左到右扫描天然满足这一点。
    """
    def emit(seg, bold):
        pos = 0
        while pos < len(seg):
            code = seg.find("`", pos)
            strong = seg.find("**", pos)
            cands = [x for x in (code, strong) if x != -1]
            if not cands:
                break
            nxt = min(cands)
            if nxt == code:
                end = seg.find("`", code + 1)
                if end == -1:
                    break
                if code > pos:
                    r = paragraph.add_run(seg[pos:code]); r.bold = bold or None
                r = paragraph.add_run(seg[code + 1:end])
                r.font.name = "Consolas"; r.bold = bold or None
                pos = end + 1
            else:
                end = seg.find("**", strong + 2)
                if end == -1:
                    break
                if strong > pos:
                    r = paragraph.add_run(seg[pos:strong]); r.bold = bold or None
                emit(seg[strong + 2:end], True)
                pos = end + 2
        if pos < len(seg):
            r = paragraph.add_run(seg[pos:]); r.bold = bold or None

    emit(text, False)


def _set_table_row_pagination(row, *, repeat_header=False):
    """Keep a row intact when possible and optionally repeat it on each page."""
    tr_pr = row._tr.get_or_add_trPr()
    if repeat_header and tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def md_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows if "---" not in r]
    if not cells:
        return
    ncol = len(cells[0])
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Light Grid Accent 1"
    # 默认是等宽固定布局，于是只放「①」的 # 列和整段正文列一样宽，
    # 正文列被挤到折 12~16 行、标识符从词中间劈开（`graph_service.upsert_skil` + `l`）。
    # 改成按内容自适应 + 占满页宽，渲染器就会把窄列压窄、把宽度让给正文列。
    t.autofit = True
    tbl_pr = t._tbl.tblPr
    for tag in ("w:tblLayout", "w:tblW"):
        for el in tbl_pr.findall(qn(tag)):
            tbl_pr.remove(el)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)
    width = OxmlElement("w:tblW"); width.set(qn("w:type"), "pct"); width.set(qn("w:w"), "5000")
    tbl_pr.append(width)
    for i, row in enumerate(cells):
        _set_table_row_pagination(t.rows[i], repeat_header=(i == 0))
        for j in range(ncol):
            cell = t.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, row[j] if j < len(row) else "")
            for run in para.runs:
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True


def convert(path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', "Microsoft YaHei")

    with open(path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                pc = doc.add_paragraph()
                r = pc.add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                # 只设 ascii/hAnsi 字体时，代码块里的中文和 ┌─┤ 制表符会回退到别的字体，
                # 字宽和 Consolas 对不上，ASCII 架构图的右边框就会参差不齐甚至跑到框外。
                # 显式把 eastAsia 也钉成等宽，整幅图才对得齐。
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                pc.paragraph_format.left_indent = Inches(0.2)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        if re.match(r"^#{1,4} ", line):
            level = len(line) - len(line.lstrip("#")) - 1   # "# "->0, "## "->1, ...
            h = doc.add_heading("", level=level)
            # 标题同样走行内解析：标题里写 `docs/测试数据/seed_jds.json` 时，
            # 直接 add_heading(text) 会把反引号原样打进 Word（§2.1/§2.2 曾如此）。
            add_runs(h, line[level + 2:].strip())
            h.paragraph_format.keep_with_next = True
            if level == 0:
                for r in h.runs:
                    r.font.color.rgb = ACCENT
        elif line.startswith(">"):
            # 收集整个引用块：连续的 `> ` 行并成一段，单独一行 `>` 表示块内换段
            quote_paras, cur = [], []
            while i < len(lines) and lines[i].startswith(">"):
                body = lines[i][1:]
                body = body[1:] if body.startswith(" ") else body
                if body.strip():
                    cur.append(body)
                else:
                    if cur:
                        quote_paras.append(cur); cur = []
                i += 1
            if cur:
                quote_paras.append(cur)
            for chunk in quote_paras:
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
                add_runs(p, join_soft_wrap(chunk))
                for r in p.runs:
                    r.italic = True; r.font.color.rgb = MUTED
            continue
        elif line.strip().startswith("|") and "|" in line:
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i]); i += 1
            md_table(doc, tbl)
            continue
        elif re.match(r"^[-*·]\s+", line.strip()) or re.match(r"^\d+\.\s+", line.strip()):
            # 缩进深度决定层级：源文件里「- **画面**：」下面的「  1. …」是子项，
            # 过去一律用 List Bullet，父项就变成一个空项目符号、子项与它平铺同级。
            indent = len(line) - len(line.lstrip(" "))
            depth = min(indent // 2, 2)          # 0 / 1 / 2 → List Bullet / 2 / 3
            txt = re.sub(r"^[-*·]\s+", "", line.strip())
            txt = re.sub(r"^\d+\.\s+", "", txt)
            # 列表项的软换行续行并进本条，否则续行会掉成一个独立段落
            cont = [txt]
            i += 1
            while i < len(lines) and not _is_block_start(lines[i]):
                cont.append(lines[i].strip()); i += 1
            p = doc.add_paragraph(style="List Bullet" if depth == 0 else f"List Bullet {depth + 1}")
            add_runs(p, join_soft_wrap(cont))
            continue
        elif line.strip() == "---":
            pass
        elif line.strip() == "":
            pass
        else:
            # 普通段落：把软换行的连续行并成一段再解析行内标记
            buf = [line.strip()]
            i += 1
            while i < len(lines) and not _is_block_start(lines[i]):
                buf.append(lines[i].strip()); i += 1
            p = doc.add_paragraph()
            add_runs(p, join_soft_wrap(buf))
            continue
        i += 1

    out = os.path.join(DOCS, os.path.splitext(os.path.basename(path))[0] + ".docx")
    doc.save(out)
    print("生成:", os.path.basename(out))


if __name__ == "__main__":
    for fn in FILES:
        convert(os.path.join(SRC, fn))
